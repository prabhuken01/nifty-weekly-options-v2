#!/usr/bin/env python3
"""
Generate PHASE2_ROADMAP.docx from markdown plan
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()
doc.add_heading('Nifty Weekly Options Strategy', 0)
doc.add_paragraph('Phase 2: AWS Deployment + Advanced Analysis + Daily Pipeline', style='Subtitle')

# CONTEXT Section
doc.add_heading('CONTEXT', level=1)

doc.add_heading('Current State', level=2)
for item in [
    'PostgreSQL database: 103,359 daily EOD records (2024-09-23 to 2026-03-24)',
    'All Greeks calculated: Delta, Gamma, Theta, Vega, Rho',
    'Database schema: option_bars_daily + option_bars_minute + data_metadata',
    'Kite API configured and tested',
    'Database not yet in cloud (still local)'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('User Goals', level=2)
for item in [
    'Push database to AWS RDS (cloud access from anywhere)',
    'Build comprehensive options analysis framework for T-1 and T expiries',
    'Set up daily incremental data fetch (Kite API + NSE Bhavcopy fallback)',
    'Create hourly scheduler for minute-level data (intra-day tracking)',
    'Generate Excel reports with analysis and probabilities'
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Expected Outcome', level=2)
p = doc.add_paragraph('Cloud-based system with daily updated Greeks and probabilities, automated hourly data refresh for live trading insights, historical analysis for strategy backtesting, and probability-based decision framework for T-1 and T expiry weeks.')

# ARCHITECTURE
doc.add_page_break()
doc.add_heading('ARCHITECTURE OVERVIEW', level=1)
doc.add_paragraph('The system flows from AWS RDS cloud database → Daily & Hourly data fetch → Options analysis → Excel reports')

# IMPLEMENTATION ROADMAP
doc.add_heading('IMPLEMENTATION ROADMAP', level=1)

# PHASE 1
doc.add_heading('PHASE 1: AWS RDS Deployment', level=2)
p = doc.add_paragraph('Files: STEP4_AWS_SETUP_GUIDE.md, backup_and_restore.py')
p.runs[0].italic = True

doc.add_heading('Steps:', level=3)
for i, item in enumerate([
    'Create AWS RDS Instance with PostgreSQL 15, db.t3.micro (free tier)',
    'Backup local database using pg_dump',
    'Restore to AWS RDS using psql command',
    'Verify AWS connection and data integrity',
    'Update all connection strings to use AWS endpoint'
], 1):
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph('Automation: Run backup_and_restore.py for one-click migration')

# PHASE 2
doc.add_heading('PHASE 2: Options Analysis Framework', level=2)
p = doc.add_paragraph('New Files: options_analysis_engine.py, probability_calculations.py, daily_analysis_report.py')
p.runs[0].italic = True

components = [
    ('A. Probability of Touching (POT)', 'Calculate % chance price touches each strike before expiry. Uses barrier option mathematics on existing Black-Scholes framework.'),
    ('B. Risk-Neutral Probability Distribution', 'Extract probability distributions from option prices. Market prices imply what probabilities traders expect.'),
    ('C. Implied vs Historical Volatility', 'Compare IV (market expectation) vs HV (realized volatility). Flag opportunities: High IV = sell, Low IV = buy.'),
    ('D. Greeks Sensitivity Analysis', 'Track daily changes in Delta, Gamma, Theta, Vega, Rho. Understand how each Greek changes with market conditions.')
]

for heading, desc in components:
    doc.add_heading(heading, level=3)
    doc.add_paragraph(desc)

# PHASE 3
doc.add_heading('PHASE 3: Daily Incremental Data Pipeline', level=2)
p = doc.add_paragraph('New File: kite_fetch_nifty_daily_incremental.py')
p.runs[0].italic = True

doc.add_paragraph('Morning Fetch (8 AM IST): Kite API (Priority 1) → NSE Bhavcopy (Priority 2) → Fallback data')

doc.add_heading('Data Sources:', level=3)
for item in [
    'Kite API: Real-time market data with authentication',
    'NSE Bhavcopy: Official end-of-day file, no API key needed',
    'Fallback: Synthetic data if both sources fail'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Cron: Daily 8 AM IST = 2:30 AM UTC')

# PHASE 4
doc.add_heading('PHASE 4: Hourly Minute-Level Data Scheduler', level=2)
p = doc.add_paragraph('Update File: kite_fetch_nifty_minute_scheduler.py')
p.runs[0].italic = True

doc.add_paragraph('Schedule: Every hour during trading (9:15 AM - 3:30 PM IST)')
doc.add_paragraph('Data: NIFTY spot + 8 contracts (T and T-1 expiry weeks)')
doc.add_paragraph('Storage: option_bars_minute table with all Greeks')

# PHASE 5
doc.add_page_break()
doc.add_heading('PHASE 5: Excel Analysis Reports', level=2)
p = doc.add_paragraph('New File: generate_excel_reports.py')
p.runs[0].italic = True

doc.add_paragraph('Report 1: Daily EOD Analysis', style='List Bullet')
eod_table = doc.add_table(rows=5, cols=2)
eod_table.style = 'Light Grid Accent 1'
headers = eod_table.rows[0].cells
headers[0].text = 'Sheet'
headers[1].text = 'Content'
eod_rows = [
    ('Today\'s EOD', 'Strike, Call/Put prices, IV, Delta, Gamma, Theta, POT, Risk-Neutral Prob'),
    ('Greeks Summary', 'Today vs Yesterday changes, Theta decay, Gamma risk visualization'),
    ('Probability Analysis', 'POT for each strike, Distribution chart, Expected move ranges'),
    ('Volatility Analysis', 'IV vs HV comparison, Volatility surface, Smile/Skew patterns')
]
for i, (sheet, content) in enumerate(eod_rows, 1):
    row_cells = eod_table.rows[i].cells
    row_cells[0].text = sheet
    row_cells[1].text = content

doc.add_paragraph()
doc.add_paragraph('Report 2: T-1 & T Expiry Strategy Report', style='List Bullet')
strategy_table = doc.add_table(rows=4, cols=2)
strategy_table.style = 'Light Grid Accent 1'
headers = strategy_table.rows[0].cells
headers[0].text = 'Sheet'
headers[1].text = 'Content'
strategy_rows = [
    ('T Expiry This Week', 'All strikes with probabilities, Greeks decay forecast, Support/Resistance'),
    ('T-1 Expiry Next Week', 'Similar analysis for next week\'s expiry, Time decay patterns'),
    ('Strategy Ideas', 'High POT strikes, Low IV opportunities, Greeks arbitrage ideas')
]
for i, (sheet, content) in enumerate(strategy_rows, 1):
    row_cells = strategy_table.rows[i].cells
    row_cells[0].text = sheet
    row_cells[1].text = content

# FILES TO CREATE
doc.add_page_break()
doc.add_heading('FILES TO CREATE/UPDATE', level=1)

doc.add_heading('NEW FILES:', level=2)
for item in [
    'options_analysis_engine.py - Core analysis orchestrator',
    'probability_calculations.py - POT, risk-neutral, volatility calculations',
    'kite_fetch_nifty_daily_incremental.py - Daily incremental data fetch',
    'kite_fetch_nifty_minute_scheduler.py - Hourly minute data fetch',
    'generate_excel_reports.py - Excel report generation',
    'setup_analysis_table.py - Create analysis_results table'
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('UPDATE FILES:', level=2)
for item in [
    'setup_schema_enhanced.py - Add analysis_results table definition',
    'README_START_HERE.md - Update with Phase 2 workflow'
]:
    doc.add_paragraph(item, style='List Bullet')

# VERIFICATION CHECKLIST
doc.add_page_break()
doc.add_heading('VERIFICATION CHECKLIST', level=1)

phases = [
    ('After Phase 1 (AWS):', [
        'AWS RDS instance created and "Available"',
        'Database backed up locally',
        'Data restored to AWS',
        'psql connects to AWS endpoint successfully',
        '103,359 records visible in AWS'
    ]),
    ('After Phase 2 (Analysis):', [
        'options_analysis_engine.py runs without errors',
        'analysis_results table populated with POT and risk-neutral probabilities',
        'Sample analysis shows POT > 0, all Greeks calculated'
    ]),
    ('After Phase 3 (Daily Pipeline):', [
        'kite_fetch_nifty_daily_incremental.py executes successfully',
        'New records appear in option_bars_daily 24 hours after first run',
        'NSE Bhavcopy fallback tested and working'
    ]),
    ('After Phase 4 (Hourly):', [
        'option_bars_minute populated hourly',
        'Cron job running at expected times',
        'Minute data timestamps match trading hours'
    ]),
    ('After Phase 5 (Excel):', [
        'Daily Excel report generated successfully',
        'All sheets contain expected data and formatting',
        'POT values range 0-1, make sense for ATM vs OTM strikes',
        'Probability distribution looks realistic'
    ])
]

for phase_heading, items in phases:
    doc.add_heading(phase_heading, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# DEPENDENCIES
doc.add_page_break()
doc.add_heading('TIMELINE & DEPENDENCIES', level=1)

doc.add_heading('Execution Order:', level=2)
phases_timeline = [
    'Phase 1 (AWS) - 15 minutes',
    'Phase 2 (Analysis) - 2-3 hours',
    'Phase 3 (Daily) - 1 hour',
    'Phase 4 (Hourly) - 30 minutes',
    'Phase 5 (Excel) - 1-2 hours'
]
for item in phases_timeline:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Python Libraries Needed:', level=2)
libs = [
    'psycopg2 (PostgreSQL) - already have',
    'pandas - already have',
    'numpy - already have',
    'scipy (for statistical distributions) - may need to install',
    'openpyxl (for Excel) - may need to install',
    'requests (for NSE download) - likely have'
]
for lib in libs:
    doc.add_paragraph(lib, style='List Bullet')

doc.add_heading('Cost Summary:', level=2)
costs = [
    'AWS Free Tier: First 12 months free',
    'After 12 months: ~10-15 USD/month for db.t3.micro RDS',
    'Data transfer: Minimal cost for within-region access'
]
for cost in costs:
    doc.add_paragraph(cost, style='List Bullet')

# Save
doc.save('PHASE2_ROADMAP.docx')
print('[OK] Document created: PHASE2_ROADMAP.docx')
